from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT_DOCX = ROOT / "output" / "docs" / "CredX_Video_Presentation_Script.docx"
LOGO = ROOT / "public" / "brand" / "credx-mark.png"

CORAL = "E4512F"
CORAL_DARK = "B93D22"
CHARCOAL = "172026"
SLATE = "53616D"
PALE = "FFF4EF"
BLUE = "EAF5FF"
GREEN = "EAF8EE"
YELLOW = "FFF8DB"
WHITE = "FFFFFF"
LINE = "D7DEE3"


def rgb(hex_color):
    return RGBColor.from_string(hex_color)


def set_cell_fill(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tcMar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def set_run(run, size=10.5, bold=False, color=CHARCOAL, font="Aptos", italic=False):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = rgb(color)


def add_text(doc, text, size=10.5, bold=False, color=CHARCOAL, after=5, before=0, align=None, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    if align is not None:
        p.alignment = align
    set_run(p.add_run(text), size=size, bold=bold, color=color, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(14 if level == 1 else 9)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run(r, size=17 if level == 1 else 13, bold=True, color=CORAL if level == 1 else CHARCOAL)
    return p


def add_callout(doc, label, text, fill=PALE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.45)
    cell = table.cell(0, 0)
    cell.width = Inches(6.45)
    set_cell_fill(cell, fill)
    set_cell_margins(cell, 140, 180, 140, 180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    set_run(p.add_run(label.upper() + "  "), size=9.5, bold=True, color=CORAL)
    set_run(p.add_run(text), size=10.5, color=CHARCOAL)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def scene(doc, number, time, page, objective, actions, script, judge):
    add_heading(doc, f"Scene {number}  |  {time}", 1)
    add_text(doc, objective, size=12, bold=True, color=CHARCOAL, after=7)
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [1.15, 5.30]
    for label, content, fill in [
        ("SHOW", page, BLUE),
        ("DO", actions, WHITE),
        ("SAY", script, PALE),
        ("JUDGE VALUE", judge, GREEN),
    ]:
        cells = table.add_row().cells
        for idx, w in enumerate(widths):
            cells[idx].width = Inches(w)
            set_cell_margins(cells[idx], 115, 130, 115, 130)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        set_cell_fill(cells[0], CORAL if label == "SAY" else CHARCOAL)
        set_cell_fill(cells[1], fill)
        p0 = cells[0].paragraphs[0]
        set_run(p0.add_run(label), size=8.5, bold=True, color=WHITE)
        p1 = cells[1].paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        p1.paragraph_format.line_spacing = 1.13
        set_run(p1.add_run(content), size=9.7, color=CHARCOAL)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def page_break(doc):
    doc.add_page_break()


def setup_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(CHARCOAL)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15
    for level in (1, 2, 3):
        s = styles[f"Heading {level}"]
        s.font.name = "Aptos Display"
        s._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        s._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        s.font.bold = True
        s.font.color.rgb = rgb(CORAL if level == 1 else CHARCOAL)


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    set_run(p.add_run("CREDX  |  VIDEO PRESENTATION PLAYBOOK"), size=8, bold=True, color=SLATE)


def build():
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_styles(doc)
    sec = doc.sections[0]
    sec.top_margin = Inches(0.72)
    sec.bottom_margin = Inches(0.68)
    sec.left_margin = Inches(0.82)
    sec.right_margin = Inches(0.82)
    sec.header_distance = Inches(0.32)
    sec.footer_distance = Inches(0.34)
    add_footer(sec)

    # Cover
    add_text(doc, "SMART JOB MATCHING DASHBOARD", size=10, bold=True, color=CORAL, after=18)
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(16)
        p.add_run().add_picture(str(LOGO), width=Inches(0.72))
    add_text(doc, "CredX Video Presentation Script", size=30, bold=True, color=CHARCOAL, after=8)
    add_text(doc, "A five-minute, judge-focused recording guide", size=16, color=SLATE, after=22)
    add_callout(doc, "Core message", "CredX delivers ranked recommendations, not just filtered results. Every match score is deterministic, explainable, and connected to an end-to-end student and recruiter workflow.", PALE)
    add_text(doc, "Recording target", size=11, bold=True, color=CORAL, before=14, after=4)
    add_text(doc, "Length: 5:00-5:30  |  Resolution: 1080p or higher  |  Browser zoom: 90%  |  Tone: confident, clear, evidence-led", size=10.5, after=15)
    add_text(doc, "How to use this guide", size=11, bold=True, color=CORAL, after=4)
    add_text(doc, "Keep this document on a second screen. For every scene, follow the four cues: SHOW the correct page, DO one visible action, SAY the narration naturally, and land the JUDGE VALUE before moving on.", size=11, after=8)
    add_text(doc, "Do not read word-for-word if it sounds unnatural. Preserve the highlighted claims and demonstrate the working feature on screen.", size=10, italic=True, color=SLATE)

    page_break(doc)
    add_heading(doc, "Run of show", 1)
    add_text(doc, "The first 20 seconds establish the problem. The middle proves end-to-end functionality. The final minute proves correctness and explainability.", size=11, color=SLATE, after=10)
    data = [
        ("0:00-0:20", "Landing page", "Problem and product promise"),
        ("0:20-0:40", "Architecture diagram", "System credibility"),
        ("0:40-0:55", "Sign-in", "Secure workflow"),
        ("0:55-1:55", "Student profile", "Structured data and resume intelligence"),
        ("1:55-2:45", "Student matches", "Ranking, score indicators, filters"),
        ("2:45-3:20", "Job details", "Explain-the-match breakdown"),
        ("3:20-3:50", "Applications", "End-to-end tracking"),
        ("3:50-4:30", "Recruiter workspace", "Bonus two-sided workflow"),
        ("4:30-5:10", "Algorithm diagram", "Correctness and scoring logic"),
        ("5:10-5:30", "Dashboard", "Memorable conclusion"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [1.15, 1.85, 3.45]
    hdr = table.rows[0].cells
    for i, (name, w) in enumerate(zip(("TIME", "SCREEN", "PURPOSE"), widths)):
        hdr[i].width = Inches(w); set_cell_fill(hdr[i], CHARCOAL); set_cell_margins(hdr[i])
        set_run(hdr[i].paragraphs[0].add_run(name), size=9, bold=True, color=WHITE)
    set_repeat_table_header(table.rows[0])
    for idx, row in enumerate(data):
        cells = table.add_row().cells
        for i, (value, w) in enumerate(zip(row, widths)):
            cells[i].width = Inches(w); set_cell_fill(cells[i], "F7F9FA" if idx % 2 else WHITE); set_cell_margins(cells[i])
            set_run(cells[i].paragraphs[0].add_run(value), size=9.3, bold=(i == 0), color=CHARCOAL)
    add_callout(doc, "Demo rule", "Never describe a feature for more than ten seconds without showing the feature or its result.", YELLOW)

    page_break(doc)
    scene(doc, "01", "0:00-0:20", "Landing page - http://localhost:3000", "Open with the problem, then state the product difference.", "Start at the hero section. Scroll only enough to reveal the product promise.", "Hello, we are presenting CredX, a smart and explainable job-matching platform for students and recruiters. Traditional portals show a large filtered list, but rarely explain which opportunities are genuinely suitable. CredX turns skills, GPA, location preferences, and work authorization into ranked recommendations with a transparent match breakdown.", "Frames the solution around ranked and explainable recommendations - the central evaluation requirement.")
    scene(doc, "02", "0:20-0:40", "CredX Production System Architecture diagram", "Establish technical credibility without over-explaining infrastructure.", "Point from users to Next.js, then MongoDB, Cloudinary, Google OAuth, and Groq.", "CredX is a modular full-stack Next.js application. Google OAuth verifies identity. MongoDB stores users, profiles, listings, matches, and applications. Cloudinary stores resumes. PDFParse, Mammoth, and Tesseract extract text, and Groq suggests skills. The match score itself remains deterministic and rule-based, so every recommendation is explainable.", "Shows architecture awareness while making clear that AI does not generate the match score.")
    scene(doc, "03", "0:40-0:55", "Custom sign-in page", "Show that protected journeys start with a real authentication workflow.", "Show the Google sign-in control. Use a pre-authenticated cut for a smooth recording.", "CredX uses Google authentication and JWT sessions. Protected student and recruiter pages require a valid session, and unauthenticated users are redirected to this sign-in experience.", "Demonstrates route protection and a complete user workflow.")

    page_break(doc)
    scene(doc, "04", "0:55-1:25", "Student > Profile", "Demonstrate the structured data used by the matching engine.", "Show skill tags, GPA, preferred location, and work authorization. Add one sample skill.", "The journey starts with a structured student profile. Students enter tagged skills, GPA, preferred location, and work-authorization status. Each field directly affects recommendations, so we avoid collecting information that has no value to the matching engine.", "Proves profile creation, usable multi-select skills, and sensible model inputs.")
    scene(doc, "05", "1:25-1:55", "Resume intelligence panel", "Demonstrate AI-assisted input without surrendering user control.", "Upload a prepared resume under 5 MB. Pause on the extracted skill suggestions, then approve or remove one.", "CredX accepts PDF, DOCX, PNG, JPEG, and WebP resumes. It validates the real file type, extracts text with PDFParse, Mammoth, or Tesseract OCR, stores the original file in Cloudinary, and asks Groq for structured skill suggestions. The student approves the final skills. AI assists profile creation; it does not control the profile or calculate the job-match score.", "Shows a robust resume workflow, graceful separation of storage and analysis, and responsible AI use.")
    scene(doc, "06", "1:55-2:05", "Profile save action", "Prove the process is automatic and end-to-end.", "Click Create profile or Save profile and wait for the success state.", "When the profile is saved, CredX automatically compares it with available listings and stores match results. No administrator or manual processing step is required.", "Directly addresses end-to-end functionality.")

    page_break(doc)
    scene(doc, "07", "2:05-2:35", "Student > Matches", "Show ranked recommendations rather than an ordinary results list.", "Point to the first three scores, matched skills, company, work mode, and sponsorship details.", "This is the main matching experience. Opportunities are sorted from highest to lowest match score, helping students focus on the strongest roles first. Every card shows the percentage and evidence behind the recommendation, including matched skills and job requirements.", "Proves sorting, visible score indicators, usability, and explainability.")
    scene(doc, "08", "2:35-2:55", "Matches filters", "Clarify the difference between filtering and scoring.", "Change work mode and sponsorship or location filters; show the list update.", "Students can refine results by location, work mode, and sponsorship. Filtering controls what is visible, while scoring controls what is most relevant. The two mechanisms work together without hiding the reason behind the ranking.", "Demonstrates filter usability and sound product reasoning.")
    scene(doc, "09", "2:55-3:20", "Student > Job details", "Make one recommendation fully understandable.", "Open a high-match role. Point to overall score, skill score, GPA fit, authorization status, and matched skills.", "Opening a role reveals the full explanation. This student matched because of skill overlap, GPA fit, and work-authorization compatibility. CredX exposes the contributing signals instead of presenting a mysterious percentage. Judges can see exactly why this student matched this listing.", "Directly satisfies the explain-the-match bonus and the problem statement's most important expectation.")

    page_break(doc)
    scene(doc, "10", "3:20-3:50", "Apply action, then Student > Applications", "Complete the student journey.", "Click Apply, then open the application tracker. Show date and status. Mention duplicate prevention.", "Applying immediately creates a tracked application. Duplicate applications are prevented, and the student can see the role, company, application date, and current status from one place. This completes the journey from profile creation to recommendation, application, and tracking.", "Shows stretch functionality and a complete end-to-end product flow.")
    scene(doc, "11", "3:50-4:30", "Recruiter dashboard > Listings > Applicant pipeline", "Demonstrate the optional recruiter-side bonus.", "Show recruiter metrics, a structured listing form, then open a listing with applicants and update one status.", "Recruiters can publish structured roles containing required skills, minimum GPA, location, work mode, and sponsorship availability. The applicant pipeline shows each candidate with profile context and match score. Recruiters can update application status, and ownership checks prevent unrelated recruiters from modifying the listing.", "Demonstrates recruiter posting, matched-candidate visibility, authorization, and two-sided platform value.")

    page_break(doc)
    scene(doc, "12", "4:30-5:10", "CredX Job Matching Algorithm diagram", "Explain the exact implemented scoring model.", "Follow the blue skills branch, purple GPA branch, yellow authorization branch, then the orange calculation and cap decision.", "The engine combines three explainable signals. Skills contribute 60 percent using Jaccard similarity: matched skills divided by all unique skills. GPA contributes 25 percent; qualified students receive 100, while near misses decay gradually across one GPA point. Work authorization contributes 15 percent. The raw score is rounded. If sponsorship is incompatible, the final score is capped at 20 so an operationally unsuitable role cannot rank highly. The score and breakdown are stored and jobs are ranked from highest to lowest.", "Proves correctness, sensible weighting, transparent rules, and the important incompatibility safeguard.")
    scene(doc, "13", "5:10-5:30", "Return to Student > Matches", "Finish with the completed journey and product differentiation.", "Pause on the ranked dashboard with at least three visibly different scores.", "To conclude, CredX delivers the complete required journey: structured profile creation, resume intelligence, automatic matching, ranked recommendations, clear explanations, filters, applications, and recruiter pipeline management. We did not build another filtered job list. We built an explainable decision experience where every recommendation has visible evidence. Thank you.", "Leaves judges with a memorable statement directly aligned to the evaluation criteria.")

    page_break(doc)
    add_heading(doc, "Matching algorithm speaker card", 1)
    add_callout(doc, "Formula", "Final raw score = round(Skill score x 0.60 + GPA score x 0.25 + Work-authorization score x 0.15)", BLUE)
    for title, body in [
        ("1. Skills - 60%", "Jaccard similarity: matched skills divided by the union of student and listing skills. If both sets are empty, the implementation returns 100 by specification."),
        ("2. GPA - 25%", "100 when GPA meets the minimum. Below the threshold, the score decays linearly from 100 to 0 across one GPA point and is clamped to the 0-100 range."),
        ("3. Work authorization - 15%", "100 when compatible; 0 when the student requires sponsorship and the employer does not offer it."),
        ("4. Compatibility cap", "When authorization is incompatible, the final score is min(raw score, 20). Near misses remain visible, but cannot appear as strong recommendations."),
        ("5. Stored explanation", "CredX stores final score, skill score, GPA score, compatibility, matched skills, and the calculation timestamp."),
    ]:
        add_text(doc, title, size=11, bold=True, color=CORAL, before=5, after=2)
        add_text(doc, body, size=10.2, after=5)
    add_callout(doc, "Never say", "Groq or AI calculates the match score. Groq only suggests resume skills; the deterministic rule-based engine calculates job compatibility.", YELLOW)

    page_break(doc)
    add_heading(doc, "Recording readiness checklist", 1)
    checks = [
        "Seed the database and confirm at least three varied listings are visible.",
        "Use a clean demonstration account and complete Google sign-in before recording.",
        "Prepare one resume under 5 MB containing React, TypeScript, SQL, Node.js, and MongoDB.",
        "Prepare one high match, one medium match, and one low or sponsorship-incompatible match.",
        "Create one existing application and set another to Under review.",
        "Confirm MongoDB, Google OAuth, Cloudinary, and Groq environment variables.",
        "Open landing, architecture, algorithm, student, and recruiter pages in separate tabs.",
        "Hide bookmarks, notifications, credentials, email addresses, and unrelated browser tabs.",
        "Record at 1080p or higher with browser zoom near 90 percent.",
        "Rehearse once with a timer; cut pauses, loading waits, and authentication consent screens.",
    ]
    for item in checks:
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        c0, c1 = table.rows[0].cells
        c0.width = Inches(0.42); c1.width = Inches(6.03)
        set_cell_fill(c0, PALE); set_cell_fill(c1, WHITE)
        set_cell_margins(c0, 90, 100, 90, 100); set_cell_margins(c1, 90, 130, 90, 130)
        p0 = c0.paragraphs[0]; p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p0.add_run("□"), size=13, bold=True, color=CORAL)
        set_run(c1.paragraphs[0].add_run(item), size=9.8, color=CHARCOAL)
    add_callout(doc, "Final confidence check", "Can a judge understand why the top job scored higher than the second job? If yes, the video has demonstrated the product's central value.", GREEN)

    # Document metadata
    props = doc.core_properties
    props.title = "CredX Video Presentation Script"
    props.subject = "Judge-focused product demo and recording playbook"
    props.author = "CredX Team"
    props.keywords = "CredX, job matching, demo script, presentation"
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build()
